import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import các hàm xử lý từ các file riêng biệt
from processor_multiple_choice import calculate_question_stats
from processor_essay import calculate_essay_stats
from processor_common import evaluate_exam_difficulty_mix

st.set_page_config(
    page_title="Exam Quality Evaluation System", 
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS cho giao diện công nghệ màu tím
st.markdown("""
<style>
    /* Main container với gradient tím */
    .main {
        background: linear-gradient(135deg, #1a1a2e 0%, #16213e 25%, #0f3460 75%, #533483 100%);
        min-height: 100vh;
    }
    
    /* Header styling */
    .stApp > header {
        background-color: transparent;
    }
    
    /* Title với gradient tím - hồng */
    h1 {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 50%, #f093fb 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: 800;
        letter-spacing: 3px;
        text-transform: uppercase;
        animation: glow 3s ease-in-out infinite alternate;
        text-shadow: 0 0 30px rgba(102, 126, 234, 0.5);
    }
    
    @keyframes glow {
        from { 
            filter: drop-shadow(0 0 20px rgba(102, 126, 234, 0.8));
        }
        to { 
            filter: drop-shadow(0 0 30px rgba(118, 75, 162, 0.8));
        }
    }
    
    /* Subheader với màu tím neon */
    h2 {
        color: #a78bfa !important;
        font-weight: 700;
        margin-top: 2rem;
        text-shadow: 0 2px 10px rgba(167, 139, 250, 0.5);
        border-bottom: 2px solid rgba(167, 139, 250, 0.3);
        padding-bottom: 0.5rem;
    }
    
    h3 {
        color: #c084fc !important;
        font-weight: 600;
        text-shadow: 0 2px 8px rgba(192, 132, 252, 0.4);
    }
    
    /* Grid layout cho columns */
    [data-testid="column"] {
        background: linear-gradient(135deg, rgba(102, 126, 234, 0.05) 0%, rgba(118, 75, 162, 0.05) 100%);
        border: 1px solid rgba(167, 139, 250, 0.2);
        border-radius: 20px;
        padding: 20px;
        backdrop-filter: blur(10px);
        box-shadow: 
            0 8px 32px 0 rgba(102, 126, 234, 0.15),
            inset 0 1px 0 0 rgba(255, 255, 255, 0.1);
        margin: 0.5rem;
    }
    
    /* Selectbox với style tím */
    .stSelectbox > div > div {
        background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
        border: 2px solid #8b5cf6;
        border-radius: 15px;
        backdrop-filter: blur(20px);
        box-shadow: 0 4px 15px rgba(139, 92, 246, 0.3);
    }
    
    /* File uploader với border animation */
    .stFileUploader > div {
        background: linear-gradient(135deg, rgba(102, 126, 234, 0.08) 0%, rgba(118, 75, 162, 0.08) 100%);
        border: 2px dashed #a78bfa;
        border-radius: 20px;
        padding: 30px;
        backdrop-filter: blur(15px);
        transition: all 0.4s cubic-bezier(0.175, 0.885, 0.32, 1.275);
        position: relative;
        overflow: hidden;
    }
    
    .stFileUploader > div:hover {
        border-color: #c084fc;
        background: linear-gradient(135deg, rgba(192, 132, 252, 0.15) 0%, rgba(167, 139, 250, 0.15) 100%);
        transform: translateY(-3px) scale(1.01);
        box-shadow: 
            0 10px 40px rgba(139, 92, 246, 0.4),
            inset 0 0 30px rgba(192, 132, 252, 0.1);
    }
    
    /* DataFrames với style tím */
    .stDataFrame {
        border-radius: 15px;
        overflow: hidden;
        box-shadow: 
            0 10px 40px rgba(102, 126, 234, 0.2),
            0 0 0 1px rgba(167, 139, 250, 0.2);
    }
    
    /* Status messages với màu phù hợp */
    .stSuccess {
        background: linear-gradient(135deg, rgba(34, 197, 94, 0.15) 0%, rgba(16, 185, 129, 0.15) 100%);
        border-left: 4px solid #10b981;
        border-radius: 12px;
        backdrop-filter: blur(10px);
    }
    
    .stInfo {
        background: linear-gradient(135deg, rgba(139, 92, 246, 0.15) 0%, rgba(167, 139, 250, 0.15) 100%);
        border-left: 4px solid #8b5cf6;
        border-radius: 12px;
        backdrop-filter: blur(10px);
    }
    
    .stWarning {
        background: linear-gradient(135deg, rgba(251, 146, 60, 0.15) 0%, rgba(250, 204, 21, 0.15) 100%);
        border-left: 4px solid #f59e0b;
        border-radius: 12px;
        backdrop-filter: blur(10px);
    }
    
    .stError {
        background: linear-gradient(135deg, rgba(239, 68, 68, 0.15) 0%, rgba(220, 38, 38, 0.15) 100%);
        border-left: 4px solid #ef4444;
        border-radius: 12px;
        backdrop-filter: blur(10px);
    }
    
    /* Button với gradient tím */
    .stDownloadButton > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        border-radius: 30px;
        padding: 12px 35px;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 2px;
        transition: all 0.3s cubic-bezier(0.175, 0.885, 0.32, 1.275);
        box-shadow: 
            0 6px 30px rgba(102, 126, 234, 0.4),
            inset 0 1px 0 rgba(255, 255, 255, 0.2);
    }
    
    .stDownloadButton > button:hover {
        background: linear-gradient(135deg, #764ba2 0%, #f093fb 100%);
        transform: translateY(-3px) scale(1.05);
        box-shadow: 
            0 10px 40px rgba(118, 75, 162, 0.5),
            inset 0 1px 0 rgba(255, 255, 255, 0.3);
    }
    
    /* Expander với style tím */
    .streamlit-expanderHeader {
        background: linear-gradient(135deg, rgba(139, 92, 246, 0.1) 0%, rgba(167, 139, 250, 0.1) 100%);
        border: 1px solid rgba(167, 139, 250, 0.3);
        border-radius: 15px;
        backdrop-filter: blur(10px);
        color: #c084fc !important;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    
    .streamlit-expanderHeader:hover {
        background: linear-gradient(135deg, rgba(139, 92, 246, 0.2) 0%, rgba(167, 139, 250, 0.2) 100%);
        border-color: #a78bfa;
    }
    
    /* Metric cards với glass effect */
    [data-testid="metric-container"] {
        background: linear-gradient(135deg, rgba(139, 92, 246, 0.1) 0%, rgba(192, 132, 252, 0.1) 100%);
        border: 1px solid rgba(167, 139, 250, 0.3);
        border-radius: 20px;
        padding: 20px;
        backdrop-filter: blur(20px);
        box-shadow: 
            0 8px 32px rgba(139, 92, 246, 0.2),
            inset 0 1px 0 rgba(255, 255, 255, 0.1);
    }
    
    /* Text color */
    p, .stMarkdown {
        color: #e9d5ff;
        line-height: 1.6;
    }
    
    /* JSON display */
    .stJson {
        background: linear-gradient(135deg, rgba(30, 27, 75, 0.8) 0%, rgba(76, 29, 149, 0.3) 100%);
        border-radius: 15px;
        padding: 20px;
        border: 1px solid rgba(167, 139, 250, 0.3);
        backdrop-filter: blur(10px);
    }
    
    /* Tables với gradient header */
    .dataframe {
        background: rgba(30, 27, 75, 0.4) !important;
        border: 1px solid rgba(167, 139, 250, 0.2);
    }
    
    .dataframe thead tr th {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
        color: white !important;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 1.5px;
        padding: 15px !important;
        border: none !important;
    }
    
    .dataframe tbody tr {
        border-bottom: 1px solid rgba(167, 139, 250, 0.1);
    }
    
    .dataframe tbody tr:hover {
        background: linear-gradient(135deg, rgba(139, 92, 246, 0.15) 0%, rgba(167, 139, 250, 0.15) 100%) !important;
        transition: all 0.3s ease;
    }
    
    /* Divider lines */
    hr {
        border: none;
        height: 2px;
        background: linear-gradient(90deg, transparent, #a78bfa, transparent);
        margin: 2rem 0;
    }
    
    /* Custom scrollbar */
    ::-webkit-scrollbar {
        width: 10px;
        height: 10px;
    }
    
    ::-webkit-scrollbar-track {
        background: rgba(30, 27, 75, 0.5);
        border-radius: 10px;
    }
    
    ::-webkit-scrollbar-thumb {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-radius: 10px;
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: linear-gradient(135deg, #764ba2 0%, #f093fb 100%);
    }
</style>
""", unsafe_allow_html=True)

# Header với animation và gradient tím
st.markdown("""
<div style="
    text-align: center; 
    margin-bottom: 3rem; 
    padding: 2rem;
    background: linear-gradient(135deg, rgba(102, 126, 234, 0.05) 0%, rgba(118, 75, 162, 0.05) 100%);
    border-radius: 30px;
    border: 1px solid rgba(167, 139, 250, 0.2);
    backdrop-filter: blur(10px);
    box-shadow: 0 10px 40px rgba(102, 126, 234, 0.1);
">
    <h1 style="font-size: 3rem; margin-bottom: 0.5rem; font-weight: 900;">
        ⚡ EXAM QUALITY EVALUATION SYSTEM
    </h1>
    <p style="
        background: linear-gradient(90deg, #a78bfa 0%, #c084fc 50%, #f0abfc 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-size: 1.3rem; 
        font-weight: 500; 
        letter-spacing: 4px; 
        text-transform: uppercase;
        margin-top: 0.5rem;
    ">
        Advanced Analytics Platform
    </p>
    <div style="
        width: 200px; 
        height: 3px; 
        background: linear-gradient(90deg, transparent, #a78bfa, #c084fc, #a78bfa, transparent); 
        margin: 25px auto; 
        animation: pulse 3s infinite;
        border-radius: 3px;
    "></div>
</div>
""", unsafe_allow_html=True)

# Container cho control panel với layout chuyên nghiệp
st.markdown("<hr style='margin: 2rem 0;'>", unsafe_allow_html=True)

with st.container():
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.markdown("""
        <div style='
            background: linear-gradient(135deg, rgba(139, 92, 246, 0.1) 0%, rgba(167, 139, 250, 0.1) 100%);
            border-radius: 25px; 
            padding: 30px; 
            border: 2px solid rgba(167, 139, 250, 0.3);
            box-shadow: 
                0 15px 35px rgba(139, 92, 246, 0.2),
                inset 0 1px 0 rgba(255, 255, 255, 0.1);
            backdrop-filter: blur(20px);
        '>
            <h3 style='
                text-align: center; 
                margin-bottom: 1.5rem;
                color: #c084fc;
                font-weight: 700;
                letter-spacing: 2px;
                text-transform: uppercase;
            '>
                ⚙️ CONTROL PANEL
            </h3>
        </div>
        """, unsafe_allow_html=True)
        
        exam_type = st.selectbox(
            "🎯 **Select Examination Type**",
            ["Trắc nghiệm", "Tự luận", "Hỗn hợp"],
            help="Choose the examination format for appropriate calculation method"
        )

# Upload section với style tím gradient
st.markdown("<hr style='margin: 2rem 0;'>", unsafe_allow_html=True)

with st.container():
    st.markdown("""
    <div style='
        text-align: center; 
        margin-bottom: 2rem;
        padding: 1.5rem;
        background: linear-gradient(135deg, rgba(102, 126, 234, 0.05) 0%, rgba(118, 75, 162, 0.05) 100%);
        border-radius: 20px;
        border: 1px solid rgba(167, 139, 250, 0.2);
    '>
        <h2 style='color: #a78bfa; margin-bottom: 0.5rem;'>📊 DATA INPUT MODULE</h2>
        <p style='color: #c4b5fd; font-size: 1rem;'>Upload your Excel file for comprehensive analysis</p>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 3, 1])
    
    with col2:
        uploaded_file = st.file_uploader(
            "**📁 Select Excel File (.xlsx)**",
            type=["xlsx"],
            help="Upload Excel file containing examination data for analysis"
        )

if uploaded_file:
    try:
        # Status badges với design mới
        st.markdown("<hr style='margin: 2rem 0;'>", unsafe_allow_html=True)
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.markdown("""
            <div style='
                background: linear-gradient(135deg, rgba(34, 197, 94, 0.2) 0%, rgba(16, 185, 129, 0.2) 100%);
                border-radius: 15px;
                padding: 15px;
                text-align: center;
                border: 1px solid rgba(34, 197, 94, 0.4);
            '>
                <p style='color: #10b981; font-weight: 700; margin: 0;'>✅ FILE STATUS</p>
                <p style='color: #86efac; margin: 5px 0 0 0;'>Loaded</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            st.markdown(f"""
            <div style='
                background: linear-gradient(135deg, rgba(139, 92, 246, 0.2) 0%, rgba(167, 139, 250, 0.2) 100%);
                border-radius: 15px;
                padding: 15px;
                text-align: center;
                border: 1px solid rgba(139, 92, 246, 0.4);
            '>
                <p style='color: #8b5cf6; font-weight: 700; margin: 0;'>🎯 EXAM TYPE</p>
                <p style='color: #c4b5fd; margin: 5px 0 0 0;'>{exam_type}</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col3:
            st.markdown("""
            <div style='
                background: linear-gradient(135deg, rgba(251, 146, 60, 0.2) 0%, rgba(250, 204, 21, 0.2) 100%);
                border-radius: 15px;
                padding: 15px;
                text-align: center;
                border: 1px solid rgba(251, 146, 60, 0.4);
            '>
                <p style='color: #f59e0b; font-weight: 700; margin: 0;'>📊 PROCESSING</p>
                <p style='color: #fde047; margin: 5px 0 0 0;'>Active</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col4:
            import datetime
            current_time = datetime.datetime.now().strftime("%H:%M:%S")
            st.markdown(f"""
            <div style='
                background: linear-gradient(135deg, rgba(59, 130, 246, 0.2) 0%, rgba(147, 51, 234, 0.2) 100%);
                border-radius: 15px;
                padding: 15px;
                text-align: center;
                border: 1px solid rgba(59, 130, 246, 0.4);
            '>
                <p style='color: #3b82f6; font-weight: 700; margin: 0;'>⏰ TIME</p>
                <p style='color: #93c5fd; margin: 5px 0 0 0;'>{current_time}</p>
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("<hr style='margin: 2rem 0;'>", unsafe_allow_html=True)
        
        # Hiển thị dữ liệu theo loại đề thi
        if exam_type == "Trắc nghiệm":
            # Đọc và hiển thị dữ liệu trắc nghiệm
            df_input = pd.read_excel(uploaded_file)
            # Chuyển đổi object về string để tránh lỗi serialization
            for col in df_input.columns:
                if df_input[col].dtype == 'object':
                    df_input[col] = df_input[col].astype(str)
            
            st.subheader("📊 Dữ liệu đã tải lên:")
            st.dataframe(df_input, use_container_width=True)
            
        elif exam_type == "Tự luận":
            # Đọc tất cả sheets cho tự luận
            excel_file = pd.ExcelFile(uploaded_file)
            sheet_names = excel_file.sheet_names
            
            df_input = pd.read_excel(uploaded_file, sheet_name=0)
            # Chuyển đổi object về string để tránh lỗi serialization
            for col in df_input.columns:
                if df_input[col].dtype == 'object':
                    df_input[col] = df_input[col].astype(str)
            
            st.subheader("📊 Dữ liệu đã tải lên:")
            if len(sheet_names) >= 2:
                # Hiển thị cả 2 sheet nếu có
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"**Sheet 1: Điểm sinh viên** ({sheet_names[0]})")
                    st.dataframe(df_input, use_container_width=True, height=300)
                    
                with col2:
                    df_max = pd.read_excel(uploaded_file, sheet_name=1)
                    st.write(f"**Sheet 2: Điểm tối đa** ({sheet_names[1]})")
                    st.dataframe(df_max, use_container_width=True, height=300)
            else:
                # Chỉ có 1 sheet
                st.dataframe(df_input, use_container_width=True)
                st.warning("⚠️ Không có sheet điểm tối đa. Sẽ sử dụng điểm cao nhất thực tế.")
                
        elif exam_type == "Hỗn hợp":
            # Đọc và hiển thị dữ liệu hỗn hợp
            excel_file = pd.ExcelFile(uploaded_file)
            sheet_names = excel_file.sheet_names
            
            if len(sheet_names) < 2:
                st.error("❌ File Excel phải có ít nhất 2 sheet: (1) Trắc nghiệm, (2) Tự luận")
            else:
                st.subheader("📊 Dữ liệu đã tải lên:")
                
                # Hiển thị 2 sheet chính
                col1, col2 = st.columns(2)
                
                df_mcq = pd.read_excel(uploaded_file, sheet_name=0)
                df_essay = pd.read_excel(uploaded_file, sheet_name=1)
                
                with col1:
                    st.write(f"**Sheet 1: Trắc nghiệm** ({sheet_names[0]})")
                    st.dataframe(df_mcq, use_container_width=True, height=300)
                
                with col2:
                    st.write(f"**Sheet 2: Tự luận** ({sheet_names[1]})")
                    st.dataframe(df_essay, use_container_width=True, height=300)
                
                # Sheet 3 nếu có
                if len(sheet_names) >= 3:
                    with st.expander("📋 Sheet 3: Điểm tối đa (nếu có)"):
                        df_max = pd.read_excel(uploaded_file, sheet_name=2)
                        st.info(f"Sheet name: {sheet_names[2]}")
                        st.dataframe(df_max, use_container_width=True)
                else:
                    st.warning("⚠️ Không có sheet điểm tối đa. Sẽ sử dụng điểm cao nhất thực tế cho tự luận.")

        # Xử lý theo hình thức đề thi
        if exam_type == "Trắc nghiệm":
            # Tính toán độ khó từng câu cho trắc nghiệm
            result_df = calculate_question_stats(df_input)

            st.subheader("📋 Kết quả tính độ khó từng câu (Trắc nghiệm):")
            st.dataframe(result_df, use_container_width=True)

            # ---- 🔹 ĐÁNH GIÁ ĐỀ THI (thêm mới) ----
            st.subheader("📊 Đánh giá tổng quan đề thi:")

            summary_df, conclusion, disc_info = evaluate_exam_difficulty_mix(
                result_df,
                tolerance=0.05,
                check_discrimination=True  # có thể bật/tắt
            )

            st.write("### 🔎 Cơ cấu độ khó so với mục tiêu")
            st.dataframe(summary_df, use_container_width=True)

            st.markdown(f"### ✅ Kết luận: **{conclusion}**")

            if disc_info:
                st.write("### 📐 Thống kê độ phân biệt")
                st.json(disc_info)


            # ---- Xuất file Word ----
            def convert_to_word(result_df, summary_df, conclusion, disc_info):
                doc = Document()

                # Tiêu đề
                title = doc.add_heading('BÁO CÁO ĐÁNH GIÁ ĐỘ KHÓ ĐỀ THI TRẮC NGHIỆM', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Kết quả từng câu
                doc.add_heading('1. Kết quả tính độ khó từng câu', level=1)

                # Thêm bảng kết quả
                table = doc.add_table(rows=1, cols=len(result_df.columns))
                table.style = 'Table Grid'

                # Header
                header_cells = table.rows[0].cells
                for i, col in enumerate(result_df.columns):
                    header_cells[i].text = str(col)

                # Data
                for _, row in result_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, val in enumerate(row):
                        row_cells[i].text = str(val)

                doc.add_page_break()

                # Đánh giá tổng quan
                doc.add_heading('2. Đánh giá tổng quan đề thi', level=1)

                # Cơ cấu độ khó
                doc.add_heading('2.1. Cơ cấu độ khó so với mục tiêu', level=2)

                # Thêm bảng summary
                summary_table = doc.add_table(rows=1, cols=len(summary_df.columns))
                summary_table.style = 'Table Grid'

                # Header
                header_cells = summary_table.rows[0].cells
                for i, col in enumerate(summary_df.columns):
                    header_cells[i].text = str(col)

                # Data
                for _, row in summary_df.iterrows():
                    row_cells = summary_table.add_row().cells
                    for i, val in enumerate(row):
                        row_cells[i].text = str(val)

                doc.add_paragraph()

                # Thống kê độ phân biệt
                if disc_info:
                    doc.add_heading('2.2. Thống kê độ phân biệt', level=2)
                    for key, value in disc_info.items():
                        doc.add_paragraph(f'{key}: {value}')

                doc.add_paragraph()

                # Kết luận
                doc.add_heading('3. Kết luận', level=1)
                conclusion_para = doc.add_paragraph(conclusion)
                conclusion_para.runs[0].bold = True

                # Lưu vào BytesIO
                output = BytesIO()
                doc.save(output)
                output.seek(0)
                return output.getvalue()


            word_data = convert_to_word(result_df, summary_df, conclusion, disc_info)
            st.download_button(
                label="⬇️ Tải báo cáo Word (.docx)",
                data=word_data,
                file_name="bao_cao_do_kho_trac_nghiem.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        elif exam_type == "Hỗn hợp":
            # Phần hiển thị đã được xử lý ở trên, bây giờ chỉ cần xử lý
            try:
                if len(sheet_names) >= 2:
                    # Lấy df_max nếu có
                    df_max = None
                    if len(sheet_names) >= 3:
                        df_max = pd.read_excel(uploaded_file, sheet_name=2)
                        
                    # Tính toán
                    from mixed_exam_evaluation import calculate_mix_stats

                    all_results = calculate_mix_stats(df_mcq, df_essay, df_max)

                st.subheader("📋 Kết quả chi tiết từng câu hỏi (Hỗn hợp):")
                st.dataframe(all_results, use_container_width=True)

                # Đánh giá tổng quan sử dụng evaluate_exam_difficulty_mix
                st.subheader("📊 Đánh giá tổng quan đề hỗn hợp:")

                # Sử dụng hàm evaluate_exam_difficulty_mix cho consistency
                summary_df, conclusion, disc_info = evaluate_exam_difficulty_mix(
                    all_results,
                    tolerance=0.05,
                    check_discrimination=True
                )

                st.write("### 🔎 Cơ cấu độ khó so với mục tiêu")
                st.dataframe(summary_df, use_container_width=True)

                st.markdown(f"### ✅ Kết luận: **{conclusion}**")

                if disc_info:
                    st.write("### 📐 Thống kê độ phân biệt")
                    st.json(disc_info)
                
                # Hiển thị thống kê riêng cho từng loại
                with st.expander("📊 Thống kê chi tiết theo loại câu hỏi"):
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.write("**Trắc nghiệm:**")
                        mc_rows = all_results[all_results['Loại câu'] == 'Trắc nghiệm']
                        if not mc_rows.empty and 'Độ khó (P)' in mc_rows.columns:
                            st.write(f"- Số câu: {len(mc_rows)}")
                            st.write(f"- Độ khó TB: {mc_rows['Độ khó (P)'].mean():.2f}")
                            if 'Độ phân biệt (D)' in mc_rows.columns:
                                st.write(f"- Độ phân biệt TB: {mc_rows['Độ phân biệt (D)'].mean():.3f}")
                    
                    with col2:
                        st.write("**Tự luận:**")
                        essay_rows = all_results[all_results['Loại câu'] == 'Tự luận']
                        if not essay_rows.empty and 'Độ khó (P)' in essay_rows.columns:
                            st.write(f"- Số câu: {len(essay_rows)}")
                            st.write(f"- Độ khó TB: {essay_rows['Độ khó (P)'].mean():.2f}")
                            if 'Độ phân biệt (D)' in essay_rows.columns:
                                st.write(f"- Độ phân biệt TB: {essay_rows['Độ phân biệt (D)'].mean():.3f}")


                # ---- Xuất file Word ----
                def convert_to_word(all_results, summary_df, conclusion, disc_info):
                    doc = Document()

                    # Tiêu đề
                    title = doc.add_heading('BÁO CÁO ĐÁNH GIÁ ĐỀ HỖN HỢP', 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Kết quả từng câu
                    doc.add_heading('1. Kết quả từng câu hỏi', level=1)
                    table = doc.add_table(rows=1, cols=len(all_results.columns))
                    table.style = 'Table Grid'
                    header_cells = table.rows[0].cells
                    for i, col in enumerate(all_results.columns):
                        header_cells[i].text = str(col)
                    for _, row in all_results.iterrows():
                        row_cells = table.add_row().cells
                        for i, val in enumerate(row):
                            row_cells[i].text = str(val)

                    doc.add_page_break()

                    # Đánh giá tổng quan
                    doc.add_heading('2. Đánh giá tổng quan đề thi', level=1)
                    
                    # Cơ cấu độ khó
                    doc.add_heading('2.1. Cơ cấu độ khó so với mục tiêu', level=2)
                    
                    # Thêm bảng summary
                    summary_table = doc.add_table(rows=1, cols=len(summary_df.columns))
                    summary_table.style = 'Table Grid'
                    
                    # Header
                    header_cells = summary_table.rows[0].cells
                    for i, col in enumerate(summary_df.columns):
                        header_cells[i].text = str(col)
                    
                    # Data
                    for _, row in summary_df.iterrows():
                        row_cells = summary_table.add_row().cells
                        for i, val in enumerate(row):
                            row_cells[i].text = str(val)
                    
                    doc.add_paragraph()
                    
                    # Thống kê độ phân biệt
                    if disc_info:
                        doc.add_heading('2.2. Thống kê độ phân biệt', level=2)
                        for key, value in disc_info.items():
                            doc.add_paragraph(f'{key}: {value}')
                    
                    doc.add_paragraph()
                    
                    # Thống kê theo loại câu
                    doc.add_heading('2.3. Thống kê theo loại câu hỏi', level=2)
                    
                    mc_rows = all_results[all_results['Loại câu'] == 'Trắc nghiệm']
                    essay_rows = all_results[all_results['Loại câu'] == 'Tự luận']
                    
                    doc.add_paragraph('Trắc nghiệm:')
                    if not mc_rows.empty and 'Độ khó (P)' in mc_rows.columns:
                        doc.add_paragraph(f'• Số câu: {len(mc_rows)}', style='List Bullet')
                        doc.add_paragraph(f'• Độ khó TB: {mc_rows["Độ khó (P)"].mean():.2f}', style='List Bullet')
                        if 'Độ phân biệt (D)' in mc_rows.columns:
                            doc.add_paragraph(f'• Độ phân biệt TB: {mc_rows["Độ phân biệt (D)"].mean():.3f}', style='List Bullet')
                    
                    doc.add_paragraph('Tự luận:')
                    if not essay_rows.empty and 'Độ khó (P)' in essay_rows.columns:
                        doc.add_paragraph(f'• Số câu: {len(essay_rows)}', style='List Bullet')
                        doc.add_paragraph(f'• Độ khó TB: {essay_rows["Độ khó (P)"].mean():.2f}', style='List Bullet')
                        if 'Độ phân biệt (D)' in essay_rows.columns:
                            doc.add_paragraph(f'• Độ phân biệt TB: {essay_rows["Độ phân biệt (D)"].mean():.3f}', style='List Bullet')
                    
                    doc.add_paragraph()
                    
                    # Kết luận
                    doc.add_heading('3. Kết luận', level=1)
                    conclusion_para = doc.add_paragraph(conclusion)
                    conclusion_para.runs[0].bold = True

                    # Lưu file
                    output = BytesIO()
                    doc.save(output)
                    output.seek(0)
                    return output.getvalue()


                word_data = convert_to_word(all_results, summary_df, conclusion, disc_info)
                st.download_button(
                        label="⬇️ Tải báo cáo Word (.docx)",
                        data=word_data,
                        file_name="bao_cao_de_hon_hop.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            except Exception as e:
                st.error(f"❌ Lỗi khi xử lý đề hỗn hợp: {e}")

        else:  # Tự luận
            st.subheader("📋 Xử lý đề thi tự luận")

            # Đọc sheet 2 nếu có (chứa điểm tối đa)
            max_scores_df = None
            try:
                # Đọc tất cả sheets
                excel_file = pd.ExcelFile(uploaded_file)
                sheet_names = excel_file.sheet_names

                if len(sheet_names) >= 2:
                    # Đọc sheet 2 (index 1)
                    max_scores_df = pd.read_excel(uploaded_file, sheet_name=sheet_names[1])
                    st.info(f"📊 Đã tìm thấy sheet điểm tối đa: {sheet_names[1]}")

                    # Hiển thị điểm tối đa
                    with st.expander("📋 Xem điểm tối đa từng câu"):
                        st.dataframe(max_scores_df, use_container_width=True)
                else:
                    st.warning("⚠️ Không tìm thấy sheet thứ 2 chứa điểm tối đa. Sẽ sử dụng điểm cao nhất thực tế.")
            except Exception as e:
                st.warning(f"⚠️ Không thể đọc sheet 2: {e}")

            # Tính toán độ khó từng câu cho tự luận
            result_df = calculate_essay_stats(df_input, max_scores_df)

            st.subheader("📋 Kết quả tính độ khó từng câu (Tự luận):")
            st.dataframe(result_df, use_container_width=True)

            # Giải thích cách tính
            with st.expander("📚 Giải thích cách tính cho tự luận"):
                st.markdown("""
                **Độ khó (P)**: 
                - Công thức: `P = (Điểm TB của tất cả SV / Điểm tối đa) × 100`
                - Điểm tối đa lấy từ sheet 2 hoặc điểm cao nhất thực tế

                **Độ phân biệt (D)**:
                - Công thức: `D = (Điểm TB nhóm cao - Điểm TB nhóm thấp) / Điểm tối đa`
                - D ≥ 0.4: Rất tốt
                - 0.3 ≤ D < 0.4: Tốt  
                - 0.2 ≤ D < 0.3: Trung bình
                - D < 0.2: Kém
                """)

            # ---- 🔹 ĐÁNH GIÁ ĐỀ THI ----
            st.subheader("📊 Đánh giá tổng quan đề thi:")

            summary_df, conclusion, disc_info = evaluate_exam_difficulty_mix(
                result_df,
                tolerance=0.05,
                check_discrimination=True
            )

            st.write("### 🔎 Cơ cấu độ khó so với mục tiêu")
            st.dataframe(summary_df, use_container_width=True)

            st.markdown(f"### ✅ Kết luận: **{conclusion}**")

            if disc_info:
                st.write("### 📐 Thống kê độ phân biệt")
                st.json(disc_info)


            # ---- Xuất file Word ----
            def convert_to_word(result_df, summary_df, conclusion, disc_info, max_scores_df=None):
                doc = Document()

                # Tiêu đề
                title = doc.add_heading('BÁO CÁO ĐÁNH GIÁ ĐỘ KHÓ ĐỀ THI TỰ LUẬN', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Kết quả từng câu
                doc.add_heading('1. Kết quả tính độ khó từng câu', level=1)

                # Thêm bảng kết quả
                table = doc.add_table(rows=1, cols=len(result_df.columns))
                table.style = 'Table Grid'

                # Header
                header_cells = table.rows[0].cells
                for i, col in enumerate(result_df.columns):
                    header_cells[i].text = str(col)

                # Data
                for _, row in result_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, val in enumerate(row):
                        row_cells[i].text = str(val)

                # Điểm tối đa nếu có
                if max_scores_df is not None:
                    doc.add_paragraph()
                    doc.add_heading('1.1. Điểm tối đa từng câu', level=2)
                    max_table = doc.add_table(rows=1, cols=len(max_scores_df.columns))
                    max_table.style = 'Table Grid'

                    # Header
                    header_cells = max_table.rows[0].cells
                    for i, col in enumerate(max_scores_df.columns):
                        header_cells[i].text = str(col)

                    # Data
                    for _, row in max_scores_df.iterrows():
                        row_cells = max_table.add_row().cells
                        for i, val in enumerate(row):
                            row_cells[i].text = str(val)

                doc.add_page_break()

                # Đánh giá tổng quan
                doc.add_heading('2. Đánh giá tổng quan đề thi', level=1)

                # Cơ cấu độ khó
                doc.add_heading('2.1. Cơ cấu độ khó so với mục tiêu', level=2)

                # Thêm bảng summary
                summary_table = doc.add_table(rows=1, cols=len(summary_df.columns))
                summary_table.style = 'Table Grid'

                # Header
                header_cells = summary_table.rows[0].cells
                for i, col in enumerate(summary_df.columns):
                    header_cells[i].text = str(col)

                # Data
                for _, row in summary_df.iterrows():
                    row_cells = summary_table.add_row().cells
                    for i, val in enumerate(row):
                        row_cells[i].text = str(val)

                doc.add_paragraph()

                # Thống kê độ phân biệt
                if disc_info:
                    doc.add_heading('2.2. Thống kê độ phân biệt', level=2)
                    for key, value in disc_info.items():
                        doc.add_paragraph(f'{key}: {value}')

                doc.add_paragraph()

                # Giải thích cách tính
                doc.add_heading('2.3. Giải thích cách tính', level=2)
                doc.add_paragraph('Độ khó (P):')
                doc.add_paragraph('• Công thức: P = (Điểm TB của tất cả SV / Điểm tối đa) × 100', style='List Bullet')
                doc.add_paragraph('• Điểm tối đa lấy từ sheet 2 hoặc điểm cao nhất thực tế', style='List Bullet')

                doc.add_paragraph('Độ phân biệt (D):')
                doc.add_paragraph('• Công thức: D = (Điểm TB nhóm cao - Điểm TB nhóm thấp) / Điểm tối đa',
                                  style='List Bullet')
                doc.add_paragraph('• D ≥ 0.4: Rất tốt', style='List Bullet')
                doc.add_paragraph('• 0.3 ≤ D < 0.4: Tốt', style='List Bullet')
                doc.add_paragraph('• 0.2 ≤ D < 0.3: Trung bình', style='List Bullet')
                doc.add_paragraph('• D < 0.2: Kém', style='List Bullet')

                doc.add_paragraph()

                # Kết luận
                doc.add_heading('3. Kết luận', level=1)
                conclusion_para = doc.add_paragraph(conclusion)
                conclusion_para.runs[0].bold = True

                # Lưu vào BytesIO
                output = BytesIO()
                doc.save(output)
                output.seek(0)
                return output.getvalue()


            word_data = convert_to_word(result_df, summary_df, conclusion, disc_info, max_scores_df)
            st.download_button(
                label="⬇️ Tải báo cáo Word (.docx)",
                data=word_data,
                file_name="bao_cao_do_kho_tu_luan.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    except Exception as e:
        st.error(f"❌ Đã xảy ra lỗi: {e}")
else:
    st.markdown("""
    <div style='
        text-align: center; 
        margin-top: 4rem; 
        padding: 60px;
        background: linear-gradient(135deg, rgba(139, 92, 246, 0.1) 0%, rgba(192, 132, 252, 0.1) 100%);
        border-radius: 30px; 
        border: 2px dashed rgba(167, 139, 250, 0.4);
        box-shadow: 
            0 20px 60px rgba(139, 92, 246, 0.15),
            inset 0 1px 0 rgba(255, 255, 255, 0.1);
        backdrop-filter: blur(10px);
    '>
        <h2 style='
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            margin-bottom: 1.5rem;
            font-size: 2.5rem;
            font-weight: 800;
        '>
            ⚡ READY TO ANALYZE
        </h2>
        <p style='color: #c4b5fd; font-size: 1.2rem; margin-bottom: 2rem;'>
            Upload your Excel file containing examination data to begin comprehensive analysis
        </p>
        <div style='
            display: inline-block;
            padding: 15px 30px;
            background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
            border-radius: 20px;
            border: 1px solid rgba(167, 139, 250, 0.3);
        '>
            <p style='color: #a78bfa; font-size: 1rem; margin: 0;'>
                📄 Supported format: <strong>.xlsx</strong>
            </p>
        </div>
        <div style='margin-top: 3rem;'>
            <p style='color: #9ca3af; font-size: 0.9rem;'>
                Drag and drop or click to browse
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)
